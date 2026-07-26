import json
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from sqlalchemy.orm import Session

from app.core.config import get_settings
from app.db.models import (
    AIDatasetExample,
    ICPProfile,
    OnboardingAnswer,
    OnboardingSession,
)


class OnboardingReportGenerator:
    def __init__(self, db: Session):
        self.db = db
        self.settings = get_settings()

    def generate(self, session_id: str) -> str:
        session = (
            self.db.query(OnboardingSession)
            .filter(OnboardingSession.id == session_id)
            .first()
        )

        if not session:
            raise ValueError(f"Session not found: {session_id}")

        answers = (
            self.db.query(OnboardingAnswer)
            .filter(OnboardingAnswer.session_id == session_id)
            .order_by(OnboardingAnswer.id.asc())
            .all()
        )

        icp = (
            self.db.query(ICPProfile)
            .filter(ICPProfile.session_id == session_id)
            .first()
        )

        dataset_examples = (
            self.db.query(AIDatasetExample)
            .filter(AIDatasetExample.session_id == session_id)
            .order_by(AIDatasetExample.id.asc())
            .all()
        )

        reports_dir = Path(self.settings.reports_dir)
        reports_dir.mkdir(parents=True, exist_ok=True)

        timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
        short_session = session_id[:8]

        filename = f"gtm_icp_review_{short_session}_{timestamp}.docx"
        report_path = reports_dir / filename

        document = Document()
        self._setup_document(document)

        self._add_title_page(document, session)
        self._add_executive_summary(document)
        self._add_user_inputs(document, answers)
        self._add_icp_output(document, icp)
        self._add_search_recipe(document, icp)
        self._add_quality_review(document, icp)
        self._add_database_summary(document, session, answers, icp, dataset_examples)
        self._add_dataset_examples(document, dataset_examples)
        self._add_client_review_notes(document)

        document.save(report_path)

        return str(report_path)

    def _setup_document(self, document: Document) -> None:
        styles = document.styles
        styles["Normal"].font.name = "Calibri"
        styles["Normal"].font.size = Pt(10.5)

        for section in document.sections:
            section.top_margin = Pt(54)
            section.bottom_margin = Pt(54)
            section.left_margin = Pt(54)
            section.right_margin = Pt(54)

    def _add_title_page(self, document: Document, session: OnboardingSession) -> None:
        title = document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = title.add_run("GTM AI ICP Builder Review Report")
        run.bold = True
        run.font.size = Pt(22)
        run.font.color.rgb = RGBColor(80, 54, 135)

        subtitle = document.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle.add_run("ICP Output + Ready Search Recipe + Evaluation Dataset Snapshot")
        subtitle_run.font.size = Pt(12)

        document.add_paragraph()

        rows = [
            ("Session ID", session.id),
            ("Mode", session.mode or ""),
            ("Status", session.status),
            ("Current Step", str(session.current_step)),
            ("Created At", str(session.created_at)),
            ("Updated At", str(session.updated_at)),
            ("Report Generated", datetime.utcnow().strftime("%Y-%m-%d %H:%M:%S UTC")),
        ]

        self._add_key_value_table(document, rows)
        document.add_page_break()

    def _add_executive_summary(self, document: Document) -> None:
        self._heading(document, "1. Executive Summary", 1)
        document.add_paragraph(
            "This document was automatically generated after the GTM AI ICP chatbot completed its flow. "
            "It captures the user's answers, the selected ICP mode, the generated ICP profile, the ready search recipe, "
            "the ICP quality review, and the tagged evaluation/fine-tuning records created in the database."
        )
        document.add_paragraph(
            "The plan recommender is intentionally kept separate from this ICP flow. "
            "This report focuses only on target profile creation and search readiness."
        )

    def _add_user_inputs(self, document: Document, answers: list[OnboardingAnswer]) -> None:
        self._heading(document, "2. User Inputs", 1)

        if not answers:
            document.add_paragraph("No onboarding answers found.")
            return

        table = document.add_table(rows=1, cols=4)
        table.style = "Table Grid"

        headers = ["#", "Question Key", "Question", "User Answer"]
        for idx, header in enumerate(headers):
            cell = table.rows[0].cells[idx]
            cell.text = header
            self._bold_cell(cell)

        for idx, answer in enumerate(answers, start=1):
            row = table.add_row().cells
            row[0].text = str(idx)
            row[1].text = answer.question_key
            row[2].text = answer.question_text
            row[3].text = answer.answer_text

    def _add_icp_output(self, document: Document, icp: ICPProfile | None) -> None:
        self._heading(document, "3. Generated ICP Profile", 1)

        if not icp:
            document.add_paragraph("No ICP profile was generated.")
            return

        rows = [
            ("Mode", icp.mode),
            ("ICP Name", icp.icp_name),
            ("ICP Summary", icp.icp_summary),
            ("Generation Method", icp.generation_method),
            ("Needs Review", icp.needs_review),
        ]

        self._add_key_value_table(document, rows)

        self._heading(document, "ICP Data", 2)
        self._add_json_block(document, icp.icp_data or {})

    def _add_search_recipe(self, document: Document, icp: ICPProfile | None) -> None:
        self._heading(document, "4. Ready Search Recipe", 1)

        if not icp:
            document.add_paragraph("No search recipe was generated.")
            return

        document.add_paragraph(
            "This search recipe is designed to be consumed by the GTM backend or reviewed by the user before running a search."
        )

        self._add_json_block(document, icp.search_recipe or {})

    def _add_quality_review(self, document: Document, icp: ICPProfile | None) -> None:
        self._heading(document, "5. ICP Quality Review", 1)

        if not icp:
            document.add_paragraph("No ICP quality review was generated.")
            return

        self._add_json_block(document, icp.icp_quality or {})

    def _add_database_summary(
        self,
        document: Document,
        session: OnboardingSession,
        answers: list[OnboardingAnswer],
        icp: ICPProfile | None,
        dataset_examples: list[AIDatasetExample],
    ) -> None:
        self._heading(document, "6. Database Records Created", 1)

        table = document.add_table(rows=1, cols=4)
        table.style = "Table Grid"

        headers = ["Database Table", "Records Created", "Purpose", "Notes"]
        for idx, header in enumerate(headers):
            cell = table.rows[0].cells[idx]
            cell.text = header
            self._bold_cell(cell)

        rows = [
            (
                "onboarding_sessions",
                "1",
                "Tracks the ICP chatbot conversation.",
                f"Session ID: {session.id}",
            ),
            (
                "onboarding_answers",
                str(len(answers)),
                "Stores each user answer.",
                "Used as input for ICP and search recipe generation.",
            ),
            (
                "icp_profiles",
                "1" if icp else "0",
                "Stores generated ICP and ready search recipe.",
                "Primary output for backend/search engine.",
            ),
            (
                "ai_dataset_examples",
                str(len(dataset_examples)),
                "Stores tagged examples for evaluation or future fine-tuning.",
                "Plan recommendation is not included in this flow.",
            ),
        ]

        for item in rows:
            row = table.add_row().cells
            for idx, value in enumerate(item):
                row[idx].text = value

    def _add_dataset_examples(self, document: Document, dataset_examples: list[AIDatasetExample]) -> None:
        self._heading(document, "7. Tagged Evaluation / Fine-Tuning Examples", 1)

        if not dataset_examples:
            document.add_paragraph("No dataset examples were created.")
            return

        for example in dataset_examples:
            self._heading(document, f"Example ID: {example.id} | Task: {example.task_type}", 2)

            rows = [
                ("Example ID", str(example.id)),
                ("Session ID", example.session_id or ""),
                ("Task Type", example.task_type),
                ("Feedback Label", example.feedback_label or "pending_review"),
                ("Feedback Notes", example.feedback_notes or ""),
                ("Created At", str(example.created_at)),
            ]

            self._add_key_value_table(document, rows)

            self._heading(document, "Tags", 3)
            self._add_json_block(document, example.tags or {})

            self._heading(document, "Input JSON", 3)
            self._add_json_block(document, example.input_json)

            self._heading(document, "Output JSON", 3)
            self._add_json_block(document, example.output_json)

    def _add_client_review_notes(self, document: Document) -> None:
        self._heading(document, "8. Client Review Notes", 1)

        table = document.add_table(rows=1, cols=4)
        table.style = "Table Grid"

        headers = ["Area", "Client Feedback", "Status", "Next Action"]
        for idx, header in enumerate(headers):
            cell = table.rows[0].cells[idx]
            cell.text = header
            self._bold_cell(cell)

        rows = [
            ("Mode Selection", "", "Pending Review", ""),
            ("Questions", "", "Pending Review", ""),
            ("ICP Quality", "", "Pending Review", ""),
            ("Search Recipe", "", "Pending Review", ""),
            ("Backend Usability", "", "Pending Review", ""),
        ]

        for item in rows:
            row = table.add_row().cells
            for idx, value in enumerate(item):
                row[idx].text = value

    def _heading(self, document: Document, text: str, level: int = 1) -> None:
        document.add_heading(text, level=level)

    def _add_key_value_table(self, document: Document, rows: list[tuple[str, Any]]) -> None:
        table = document.add_table(rows=1, cols=2)
        table.style = "Table Grid"

        header_cells = table.rows[0].cells
        header_cells[0].text = "Field"
        header_cells[1].text = "Value"
        self._bold_cell(header_cells[0])
        self._bold_cell(header_cells[1])

        for key, value in rows:
            row = table.add_row().cells
            row[0].text = str(key)
            row[1].text = "" if value is None else str(value)

    def _add_json_block(self, document: Document, data: Any) -> None:
        text = json.dumps(data, indent=2, ensure_ascii=False)

        paragraph = document.add_paragraph()
        run = paragraph.add_run(text)
        run.font.name = "Courier New"
        run.font.size = Pt(8)

    def _bold_cell(self, cell) -> None:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True